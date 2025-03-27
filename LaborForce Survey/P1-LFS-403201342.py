# بخش اول :رسم نمودار خطی متغیر های خواسته شده 


import os
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches

# مسیر پوشه حاوی فایل‌های پاکسازی‌شده
folder_path = r"C:\Users\Asus\Desktop\ترم2اقتصاد\AppliedEconometrics\FirstProject\Q5"

# سال‌های مورد نظر
years = range(1396, 1403)
results = []

for year in years:
    file_path = os.path.join(folder_path, f"LFS{year}_cleaned.dta")
    try:
        df = pd.read_stata(file_path)
    except Exception as e:
        print(f"Error reading file {file_path}: {e}")
        continue

    df['Age'] = pd.to_numeric(df['Age'], errors='coerce')
    df = df.dropna(subset=['Age'])
    df_filtered = df[(df['Age'] >= 15) & (df['Age'] <= 65)]

    groups = {
        'All': df_filtered,
        'Men': df_filtered[df_filtered['Sex'] == "1"],
        'Women': df_filtered[df_filtered['Sex'] == "2"]
    }
    
    for group_name, group_df in groups.items():
        total_weight = group_df['Weight'].sum()
        employed_weight = group_df[group_df['ActivityStatus'] == "1"]['Weight'].sum()
        unemployed_weight = group_df[group_df['ActivityStatus'] == "2"]['Weight'].sum()
        labor_force_weight = employed_weight + unemployed_weight  

        employment_rate_weighted = employed_weight / total_weight * 100 if total_weight > 0 else None
        unemployment_rate_weighted = unemployed_weight / labor_force_weight * 100 if labor_force_weight > 0 else None
        
        results.append({
            'Year': year,
            'Group': group_name,
            'Weighted Employment Rate': employment_rate_weighted,
            'Weighted Unemployment Rate': unemployment_rate_weighted,
            'Weighted Employed Ratio': employment_rate_weighted
        })

# ایجاد DataFrame نتایج
result_df = pd.DataFrame(results)

# ایجاد سند Word
doc = Document()
doc.add_heading('Employment Analysis Report', level=1)

# اضافه کردن جدول نتایج به Word
doc.add_heading('Weighted Employment and Unemployment Rates', level=2)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'

# اضافه کردن هدرها
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Year'
hdr_cells[1].text = 'Group'
hdr_cells[2].text = 'Employment Rate (%)'
hdr_cells[3].text = 'Unemployment Rate (%)'

# پر کردن جدول با داده‌ها
for _, row in result_df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(row['Year'])
    row_cells[1].text = row['Group']
    row_cells[2].text = f"{row['Weighted Employment Rate']:.2f}" if row['Weighted Employment Rate'] is not None else "N/A"
    row_cells[3].text = f"{row['Weighted Unemployment Rate']:.2f}" if row['Weighted Unemployment Rate'] is not None else "N/A"

# تنظیم اندازه کل شکل
fig, axs = plt.subplots(2 , 2, figsize=(14, 10))
plt.rcParams.update({'font.size': 12})

# نمودار 1: Weighted Employment Rate
sns.lineplot(data=result_df, x='Year', y='Weighted Employment Rate', hue='Group', marker='o', ax=axs[0, 0], legend=False)
axs[0, 0].set_xlabel("Year", fontsize=10)
axs[0, 0].set_ylabel("Weighted Employment Rate (%)", fontsize=8)
axs[0, 0].set_title("Weighted Employment Rate by Year and Group", fontsize=8)
axs[0, 0].grid(True)

# نمودار 2: Weighted Unemployment Rate
sns.lineplot(data=result_df, x='Year', y='Weighted Unemployment Rate', hue='Group', marker='o', ax=axs[0, 1], legend=False)
axs[0, 1].set_xlabel("Year", fontsize=10)
axs[0, 1].set_ylabel("Weighted Unemployment Rate (%)", fontsize=8)
axs[0, 1].set_title("Weighted Unemployment Rate by Year and Group", fontsize=8)
axs[0, 1].grid(True)

# نمودار 3: Weighted Employed Ratio to Total Population
sns.lineplot(data=result_df, x='Year', y='Weighted Employed Ratio', hue='Group', marker='o', ax=axs[1, 0])
axs[1, 0].set_xlabel("Year", fontsize=10)
axs[1, 0].set_ylabel("Weighted Employed Ratio to Total Population (%)", fontsize=8)
axs[1, 0].set_title("Weighted Employed Ratio to Total Population by Year and Group", fontsize=8)
axs[1, 0].grid(True)

axs[1, 0].legend(title="Group", fontsize=8, loc='upper left', bbox_to_anchor=(1.05, 1))

# حذف subplot اضافی
fig.delaxes(axs[1, 1])

# ذخیره نمودار به عنوان تصویر
chart_path = "employment_trends.png"
plt.savefig(chart_path, bbox_inches='tight', dpi=300)
plt.close()

# اضافه کردن تصویر نمودار به Word
doc.add_page_break()
doc.add_heading('Employment Trends Over Time', level=2)
doc.add_picture(chart_path, width=Inches(6))
doc.add_paragraph("The above figure illustrates the employment and unemployment trends over the years for different groups.")

# ذخیره فایل Word
word_file_path = "Employment_Analysis_Report.docx"
doc.save(word_file_path)

print(f"فایل Word با نام '{word_file_path}' ایجاد شد و شامل نتایج و نمودارها است.")
# بخش دوم رسم نقشه ایران به همراه محاسبه مقادیر
import geopandas as gpd
import matplotlib.pyplot as plt
import os

# 📌 مسیر فایل `shapefile` استان‌های ایران
shapefile_path = r"C:\Users\Asus\Desktop\ترم2اقتصاد\AppliedEconometrics\FirstProject\Q5\Province\gadm41_IRN_2.shp"

# ✅ 1. خواندن نقشه ایران
iran_map = gpd.read_file(shapefile_path)

# ✅ 2. نگه داشتن فقط استان‌ها (بدون شهرستان‌ها)
iran_provinces = iran_map.dissolve(by="NAME_1").reset_index()

# ✅ 3. نگاشت نام استان‌ها برای تطبیق با داده‌های آماری
province_name_mapping = {
    "Alborz": "Alborz",
    "Ardebil": "Ardabil",
    "Bushehr": "Bushehr",
    "Chahar Mahall and Bakhtiari": "Chaharmahal and Bakhtiari",
    "East Azarbaijan": "East Azerbaijan",
    "Esfahan": "Isfahan",
    "Fars": "Fars",
    "Gilan": "Gilan",
    "Golestan": "Golestan",
    "Hamadan": "Hamedan",
    "Hormozgan": "Hormozgan",
    "Ilam": "Ilam",
    "Kerman": "Kerman",
    "Kermanshah": "Kermanshah",
    "Khuzestan": "Khuzestan",
    "Kohgiluyeh and Buyer Ahmad": "Kohgiluyeh and Boyer-Ahmad",
    "Kordestan": "Kurdistan",
    "Lorestan": "Lorestan",
    "Markazi": "Markazi",
    "Mazandaran": "Mazandaran",
    "North Khorasan": "North Khorasan",
    "Qazvin": "Qazvin",
    "Qom": "Qom",
    "Razavi Khorasan": "Razavi Khorasan",
    "Semnan": "Semnan",
    "Sistan and Baluchestan": "Sistan and Baluchestan",
    "South Khorasan": "South Khorasan",
    "Tehran": "Tehran",
    "West Azarbaijan": "West Azerbaijan",
    "Yazd": "Yazd",
    "Zanjan": "Zanjan"
}

iran_provinces["Province Name"] = iran_provinces["NAME_1"].map(province_name_mapping)

# ✅ 4. داده‌های درصد بیمه شاغلین (خروجی جدول قبلی)
insurance_data = {
    "Kerman": 24.5, "West Azerbaijan": 26.0, "Lorestan": 26.6, "Kurdistan": 29.3,
    "Sistan and Baluchestan": 29.6, "Razavi Khorasan": 31.0, "North Khorasan": 31.0,
    "Ardabil": 33.7, "Kermanshah": 34.6, "Hamedan": 36.6, "Gilan": 36.7,
    "Golestan": 36.7, "Ilam": 37.2, "Khuzestan": 37.8, "Mazandaran": 37.9,
    "Hormozgan": 40.3, "East Azerbaijan": 41.0, "Zanjan": 43.8, "Qom": 48.5,
    "Fars": 49.3, "Kohgiluyeh and Boyer-Ahmad": 50.4, "Isfahan": 50.7,
    "Tehran": 50.8, "Qazvin": 52.3, "Alborz": 52.9, "Chaharmahal and Bakhtiari": 54.1,
    "Bushehr": 54.3, "South Khorasan": 56.1, "Markazi": 62.1, "Yazd": 62.6,
    "Semnan": 67.4
}

# ✅ 5. اضافه کردن داده‌های بیمه به `iran_provinces`
iran_provinces["Insurance Coverage (%)"] = iran_provinces["Province Name"].map(insurance_data)

# 📌 **رفع مشکل نمایش استان‌های بدون مقدار بیمه**
iran_provinces["Insurance Coverage (%)"] = iran_provinces["Insurance Coverage (%)"].fillna(0)

# ✅ 6. رسم نقشه ایران با **درصد بیمه شاغلین**
fig, ax = plt.subplots(figsize=(12, 12))

# 📌 استفاده از رنگ‌بندی بر اساس درصد بیمه‌شدگان
iran_provinces.plot(column="Insurance Coverage (%)", cmap="coolwarm_r", linewidth=0.8, edgecolor="black", ax=ax, legend=True)

# 📌 تنظیمات گرافیکی
ax.set_title("Insurance Coverage of Employed Individuals in Iran - 1402", fontsize=14, fontweight="bold")
ax.axis("off")  # حذف محورهای مختصات

# 📌 افزودن **اعداد درصد بیمه** به‌جای نام استان‌ها
for idx, row in iran_provinces.iterrows():
    if row["Insurance Coverage (%)"] > 0:
        plt.annotate(text=f"{row['Insurance Coverage (%)']:.1f}%", 
                     xy=row["geometry"].representative_point().coords[0], 
                     ha='center', fontsize=8, color="black", fontweight="bold")

# 📌 **ذخیره نمودار در همان دایرکتوری که فایل `.shp` در آن قرار دارد**
save_directory = os.path.dirname(shapefile_path)  # پوشه‌ای که فایل داده در آن قرار دارد
save_path = os.path.join(save_directory, "Insurance_Coverage_Map.png")  # مسیر ذخیره نمودار
plt.savefig(save_path, dpi=300, bbox_inches="tight")

print(f"✅ نقشه ذخیره شد: {save_path}")

# 📌 نمایش نقشه
plt.show()
# بخش سوم محاسبه درصد ها اشتغال در بخش های مختلف به تفکیک خواسته شده
import pandas as pd
import os
import matplotlib.pyplot as plt

# 📌 مسیر فایل داده
file_path = r"C:\Users\Asus\Desktop\ترم2اقتصاد\AppliedEconometrics\FirstProject\Q5\LFS1402_cleaned.dta"

# 📌 خواندن داده‌ها
df = pd.read_stata(file_path)

# 📌 تبدیل `ActivityStatus`, `Mozd_self`, و `ISIC` به عددی
df["ActivityStatus"] = pd.to_numeric(df["ActivityStatus"], errors="coerce")
df["Mozd_self"] = pd.to_numeric(df["Mozd_self"], errors="coerce")  # نوع اشتغال
df["ISIC"] = pd.to_numeric(df["ISIC"], errors="coerce")  # بخش فعالیت اقتصادی

# 📌 حذف ردیف‌هایی که مقدار `ActivityStatus` آنها خالی است
df = df.dropna(subset=["ActivityStatus", "Mozd_self", "ISIC"])

# 📌 فیلتر کردن فقط افراد شاغل (`ActivityStatus == 1`)
df_employed = df[df["ActivityStatus"] == 1]

# ✅ تعریف دسته‌بندی بخش‌های اقتصادی بر اساس `ISIC`
def classify_sector(isic):
    if 1 <= isic <= 3:
        return "Agriculture"
    elif 5 <= isic <= 43:
        return "Industry"
    elif 45 <= isic <= 99:
        return "Services"
    else:
        return "Other"

df_employed["Sector"] = df_employed["ISIC"].apply(classify_sector)

# ✅ تعریف دسته‌بندی نوع اشتغال از `Mozd_self`
def classify_employment(mozd_self):
    if mozd_self in [1, 2]:
        return "Self-Employed"
    elif mozd_self in [4, 5, 6]:
        return "Wage Earner"
    else:
        return "Other"

df_employed["Employment Type"] = df_employed["Mozd_self"].apply(classify_employment)

# 📌 محاسبه تعداد و درصد هر نوع اشتغال در هر بخش اقتصادی
sector_stats = df_employed.groupby(["Sector", "Employment Type"]).apply(lambda group: {
    "Total Weight": group["Weight"].sum()
}).apply(pd.Series)

# 📌 محاسبه درصد هر دسته نسبت به کل شاغلین در همان بخش
sector_totals = df_employed.groupby("Sector")["Weight"].sum().rename("Sector Total")
sector_stats = sector_stats.join(sector_totals, on="Sector")
sector_stats["Percentage (%)"] = (sector_stats["Total Weight"] / sector_stats["Sector Total"]) * 100

# 📌 تبدیل داده‌ها برای رسم نمودار
pivot_table = sector_stats.pivot_table(values="Percentage (%)", index="Sector", columns="Employment Type", fill_value=0)

# 📌 رسم نمودار
fig, ax = plt.subplots(figsize=(8, 5))  # تنظیم اندازه نمودار

pivot_table.plot(kind="bar", stacked=True, ax=ax, colormap="viridis")

plt.xlabel("Economic Sector")
plt.ylabel("Percentage of Employment (%)")
plt.title("Employment Type Distribution in Economic Sectors (1402)")
plt.grid(axis="y", linestyle="--", alpha=0.7)

# 📌 افزایش فاصله پایین نمودار و جابه‌جایی آن به سمت چپ
fig.subplots_adjust(left=0.15, bottom=0.25, right=0.7)  # تنظیم فاصله‌های نمودار

# 📌 تنظیم legend در سمت راست نمودار
legend = plt.legend(title="Employment Type", bbox_to_anchor=(1.05, 1), loc="upper left")
legend.get_frame().set_alpha(0.9)  # شفاف کردن پس‌زمینه‌ی legend

# 📌 ذخیره نمودار در همان دایرکتوری فایل داده
save_path = os.path.join(os.path.dirname(file_path), "Employment_Type_Distribution.png")
plt.savefig(save_path, dpi=300, bbox_inches="tight")

print(f"✅ نمودار ذخیره شد: {save_path}")

# 📌 نمایش نمودار
plt.show()
# بخش چهارم مربوط به ISCO
import pandas as pd
import os
import matplotlib.pyplot as plt

# 📌 مسیر پوشه فایل‌ها
folder_path = r"C:\Users\Asus\Desktop\ترم2اقتصاد\AppliedEconometrics\FirstProject\Q5"

# 📌 سال‌های مورد بررسی
years = list(range(1396, 1403))

# 📌 ایجاد یک دیتافریم خالی برای نگهداری داده‌های همه سال‌ها
all_data = pd.DataFrame()

# ✅ خواندن داده‌های هر سال و ترکیب آن‌ها
for year in years:
    file_path = os.path.join(folder_path, f"LFS{year}_cleaned.dta")
    
    if os.path.exists(file_path):
        df = pd.read_stata(file_path)
        df["Year"] = year
        all_data = pd.concat([all_data, df], ignore_index=True)
    else:
        print(f"⚠ فایل {file_path} یافت نشد!")

# 📌 حذف مقادیر `NaN` در `Year`, `ISCO`, و `Weight`
all_data = all_data.dropna(subset=["Year", "ISCO", "Weight"])

# ✅ تبدیل `ISCO` به رشته (`str`) برای حفظ صفرهای ابتدایی
all_data["ISCO"] = all_data["ISCO"].astype(str)

# ✅ تابع دسته‌بندی مشاغل بر اساس **رقم اول `ISCO`**
def classify_isco(isco_code):
    if pd.isna(isco_code) or len(isco_code) == 0:
        return "Unknown"
    
    first_digit = isco_code[0]

    if first_digit == "1":
        return "Legislators & Senior Managers"
    elif first_digit == "2":
        return "Professionals"
    elif first_digit == "3":
        return "Technicians & Assistants"
    elif first_digit in ["4", "5"]:
        return "Service Workers & Salespeople"
    elif first_digit in ["6", "7", "8"]:
        return "Craft & Related Workers"
    elif first_digit == "9":
        return "Elementary Workers"
    else:
        return "Other"

# 📌 اعمال دسته‌بندی به داده‌ها
all_data["Job Category"] = all_data["ISCO"].apply(classify_isco)

# 📌 محاسبه وزن کل نیروی کار در هر سال
total_weight_per_year = all_data.groupby("Year")["Weight"].sum().rename("Total Yearly Weight")

# 📌 محاسبه میانگین وزن‌شده افراد در هر گروه شغلی برای هر سال
job_stats_by_year = all_data.groupby(["Year", "Job Category"])["Weight"].sum().reset_index()

# 📌 الحاق وزن کل سالانه به داده‌های گروه‌بندی‌شده
job_stats_by_year = job_stats_by_year.merge(total_weight_per_year, on="Year")

# 📌 نرمال‌سازی: محاسبه سهم هر گروه از کل وزن نیروی کار آن سال
job_stats_by_year["Normalized Weight (%)"] = (job_stats_by_year["Weight"] / job_stats_by_year["Total Yearly Weight"]) * 100

# 📌 تبدیل داده‌ها به فرمت جدول مناسب برای رسم نمودار
pivot_table = job_stats_by_year.pivot(index="Year", columns="Job Category", values="Normalized Weight (%)").fillna(0)

# 📌 رسم نمودار
fig, ax = plt.subplots(figsize=(10, 6))  # کوچکتر کردن عرض نمودار

pivot_table.plot(kind="bar", stacked=True, colormap="viridis", ax=ax)

plt.xlabel("Year")
plt.ylabel("Normalized Employment Share (%)")
plt.title("Employment Share by Job Category (1396-1402)")
plt.grid(axis="y", linestyle="--", alpha=0.7)

# 📌 کوچک کردن نمودار و تنظیم legend در سمت راست
fig.subplots_adjust(right=0.7)  # افزایش فضای سمت راست برای legend

legend = plt.legend(title="Job Category", bbox_to_anchor=(1.05, 1), loc="upper left", frameon=True)
legend.get_frame().set_alpha(0.9)  # پس‌زمینه‌ی کمی شفاف برای خوانایی بهتر

# 📌 ذخیره نمودار در همان دایرکتوری
save_path = os.path.join(folder_path, "Employment_Share.png")
plt.savefig(save_path, dpi=300, bbox_inches="tight")  # کیفیت بالا

print(f"✅ نمودار با موفقیت در مسیر زیر ذخیره شد:\n{save_path}")

plt.show()
from docx import Document
from docx.shared import Inches
import os
import pandas as pd

# 📌 مسیر دسکتاپ برای ذخیره فایل Word
desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
word_file_path = os.path.join(desktop_path, "Employment_Analysis_Report.docx")

# ✅ ایجاد سند Word
doc = Document()
doc.add_heading('Employment Analysis Report', level=1)

# 📌 بررسی اینکه آیا `result_df` در محیط موجود است
if 'result_df' in locals():
    # 📌 اضافه کردن جدول داده‌ها
    doc.add_heading('Weighted Employment and Unemployment Rates', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # اضافه کردن هدرها
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Year'
    hdr_cells[1].text = 'Group'
    hdr_cells[2].text = 'Employment Rate (%)'
    hdr_cells[3].text = 'Unemployment Rate (%)'

    # پر کردن جدول با داده‌ها
    for _, row in result_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Year'])
        row_cells[1].text = row['Group']
        row_cells[2].text = f"{row['Weighted Employment Rate']:.2f}" if row['Weighted Employment Rate'] is not None else "N/A"
        row_cells[3].text = f"{row['Weighted Unemployment Rate']:.2f}" if row['Weighted Unemployment Rate'] is not None else "N/A"

else:
    doc.add_paragraph("❌ No employment data found!")

# 📌 اضافه کردن تصاویر نمودارها به Word
doc.add_page_break()
doc.add_heading('Employment Trends Over Time', level=2)

# 📌 مسیر نمودارهای ذخیره‌شده
chart_paths = [
    os.path.join(folder_path, "employment_trends.png"),
    os.path.join(os.path.dirname(shapefile_path), "Insurance_Coverage_Map.png"),
    os.path.join(os.path.dirname(file_path), "Employment_Type_Distribution.png"),
    os.path.join(folder_path, "Employment_Share.png")
]

# 📌 اضافه کردن نمودارها به سند
for chart_path in chart_paths:
    if os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(6))
        doc.add_paragraph(f"Figure: {os.path.basename(chart_path).replace('_', ' ').replace('.png', '')}")
        doc.add_page_break()
    else:
        doc.add_paragraph(f"❌ Image not found: {chart_path}")

# 📌 ذخیره فایل Word در دسکتاپ
doc.save(word_file_path)
print(f"✅ گزارش نهایی در مسیر زیر ذخیره شد:\n{word_file_path}")
